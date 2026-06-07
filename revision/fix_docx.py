import docx
import re

def replace_text_in_paragraph(p, replacements):
    text = p.text
    changed = False
    for old, new in replacements:
        if old in text:
            text = text.replace(old, new)
            changed = True
    
    if changed:
        # Save the first run's style (or just use default)
        style = p.style
        p.clear()
        p.add_run(text)
        p.style = style
        return True
    return False

def main():
    doc_path = r'c:\Users\MM\Desktop\nature isa\submission ready\manuscript_Brain.docx'
    doc = docx.Document(doc_path)
    
    replacements = [
        ("We further validate the model mechanistically using a generative biophysical ordinary differential equation (ODE) network-diffusion model",
         "We illustrate the potential mechanism of this synergy using a biophysical ordinary differential equation (ODE) network-diffusion model"),
         
        ("The interaction effect was robustly confirmed across three independent cohorts, validating its generalisability",
         "The interaction effect was confirmed in the primary PPMI and ADNI cohorts, with partial replication observed in OASIS-3, suggesting generalisability while highlighting cohort-specific variance"),
         
        ("The negative coefficient in PPMI (β = −0.070)",
         "The negative coefficient in PPMI (β = −0.033)"),
         
        ("Biophysical modelling confirms mechanistic origin of synergistic atrophy",
         "Biophysical modelling illustrates potential mechanistic origin of synergistic atrophy"),
         
        ("accurately reproduced the empirical pattern of neurodegeneration",
         "reproduced the broad spatial distribution of empirical neurodegeneration"),
         
        ("confirming that the observed atrophy topography emerges",
         "demonstrating that the observed atrophy topography could emerge"),
         
        ("Kaplan-Meier analysis revealed that Vulnerable Hubs subjects had a significantly higher probability of cognitive decline during follow-up (log-rank p\u2009=\u20090.026; Figure 5B), demonstrating that the network-genetic synergy score carries independent prognostic information beyond standard clinical variables.",
         "Kaplan-Meier analysis for cognitive decline (MoCA drop \u2265 3 points) did not show a statistically significant difference between groups (log-rank p = 0.51, bootstrap 95% CI [0.00, 6.80]). These findings indicate that while network-genetic synergy robustly structures group-average atrophy patterns, the current subject-level formulation does not carry independent prognostic value for individual outcomes beyond standard clinical variables."),
         
        ("log-rank p\u2009=\u20090.026)",
         "log-rank p\u2009=\u20090.51)"),
         
        ("A larger OASIS-3 or a pooled cross-cohort analysis would provide more definitive replication in this dataset.",
         "A larger OASIS-3 or a pooled cross-cohort analysis would provide more definitive replication in this dataset. Furthermore, our structural MRI measurements of perivascular spaces do not explicitly differentiate between pathologically enlarged and normal PVS, and they are susceptible to misclassification in standard T1-weighted images without paired T2-weighted scans.")
    ]
    
    count = 0
    for p in doc.paragraphs:
        if replace_text_in_paragraph(p, replacements):
            count += 1
            
    print(f"Modified {count} paragraphs.")
    doc.save(r'c:\Users\MM\Desktop\nature isa\submission ready\manuscript_Brain_revised.docx')

if __name__ == '__main__':
    main()
